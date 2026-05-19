#!/usr/bin/env python3
"""Convierte un documento DOCX del Informe Anual a PDF usando python-docx + fpdf2.

Uso:
    python convertir_docx_a_pdf.py <input.docx> [output.pdf]

Si no se especifica output.pdf, se genera en el mismo directorio que el input
con extension .pdf.
"""

import sys
import os
from pathlib import Path

from docx import Document
from fpdf import FPDF


# Ruta a fuente Unicode TrueType
FONT_DIR = os.path.join(os.path.dirname(__file__), '_fonts')
FONT_REGULAR = os.path.join(FONT_DIR, 'DejaVuSans.ttf')
FONT_BOLD = os.path.join(FONT_DIR, 'DejaVuSans-Bold.ttf')
FONT_ITALIC = os.path.join(FONT_DIR, 'DejaVuSans-Oblique.ttf')
FONT_BI = os.path.join(FONT_DIR, 'DejaVuSans-BoldOblique.ttf')

# Fallback: fuentes del sistema en Windows
WINDOWS_FONTS = r'C:\Windows\Fonts'
FALLBACK_REGULAR = os.path.join(WINDOWS_FONTS, 'arial.ttf')
FALLBACK_BOLD = os.path.join(WINDOWS_FONTS, 'arialbd.ttf')
FALLBACK_ITALIC = os.path.join(WINDOWS_FONTS, 'ariali.ttf')
FALLBACK_BI = os.path.join(WINDOWS_FONTS, 'arialbi.ttf')


def _find_font(primary, secondary):
    """Busca fuente, probando primary luego secondary."""
    for path in (primary, secondary):
        if os.path.exists(path):
            return path
    # Último recurso: cualquier .ttf en el directorio _fonts
    if os.path.isdir(FONT_DIR):
        for f in os.listdir(FONT_DIR):
            if f.endswith('.ttf'):
                return os.path.join(FONT_DIR, f)
    return None


class InformePDF(FPDF):
    """PDF personalizado con Unicode y formato profesional."""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._setup_fonts()

    def _setup_fonts(self):
        """Registra fuentes Unicode (TTF)."""
        r = _find_font(FONT_REGULAR, FALLBACK_REGULAR)
        b = _find_font(FONT_BOLD, FALLBACK_BOLD)
        i = _find_font(FONT_ITALIC, FALLBACK_ITALIC)
        bi = _find_font(FONT_BI, FALLBACK_BI)

        if not r:
            raise RuntimeError(
                "No se encontro ninguna fuente TTF. "
                "Descarga DejaVu Sans desde https://dejavu-fonts.github.io/ "
                "y coloca los .ttf en '_fonts/' junto a este script."
            )

        self.add_font('Custom', '', r, uni=True)
        self.add_font('Custom', 'B', b or r, uni=True)
        self.add_font('Custom', 'I', i or r, uni=True)
        self.add_font('Custom', 'BI', bi or r, uni=True)

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font('Custom', 'I', 7)
        self.set_text_color(0x99, 0x99, 0x99)
        self.cell(0, 5, 'Hotel Posada de la Silleria - Informe Anual 2025', align='C')
        self.ln(8)

    def footer(self):
        if self.page_no() == 1:
            return
        self.set_y(-15)
        self.set_font('Custom', 'I', 7)
        self.set_text_color(0x99, 0x99, 0x99)
        self.cell(0, 10, f'Pagina {self.page_no() - 1}', align='C')


def docx_a_pdf(docx_path, pdf_path=None):
    """Convierte un DOCX del informe anual a PDF."""
    if pdf_path is None:
        pdf_path = Path(docx_path).with_suffix('.pdf')

    doc = Document(docx_path)

    pdf = InformePDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Colores
    AZUL_OSCURO = (27, 58, 92)
    AZUL_MEDIO = (74, 111, 140)
    GRIS_TEXTO = (60, 60, 60)
    GRIS_CLARO = (200, 200, 200)

    page_w = 190  # ancho util A4 con margen 10mm c/u
    margen_izq = 10

    pdf.set_margins(margen_izq, 15, 10)

    # ── Portada ──
    pdf.ln(60)
    pdf.set_font('Custom', 'B', 28)
    pdf.set_text_color(*AZUL_OSCURO)
    pdf.cell(page_w, 14, 'INFORME ANUAL 2025', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(20)

    pdf.set_font('Custom', 'B', 18)
    pdf.set_text_color(*AZUL_MEDIO)
    pdf.cell(page_w, 11, 'Hotel Posada de la Silleria', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(16)

    pdf.set_font('Custom', '', 14)
    pdf.set_text_color(0x66, 0x66, 0x66)
    pdf.cell(page_w, 9, 'Toledo, Espana', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(12)

    pdf.set_font('Custom', '', 11)
    pdf.set_text_color(0x99, 0x99, 0x99)
    pdf.cell(page_w, 8, 'Generado: 15 de mayo de 2026', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(8)

    pdf.set_font('Custom', 'I', 9)
    pdf.cell(page_w, 7, 'Datos sinteticos (seed=42) basados en motor de revenue management', align='C', new_x='LMARGIN', new_y='NEXT')

    pdf.add_page()

    # ── Procesar parrafos ──
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style else ''

        # Saltar parrafos vacios de la portada
        if not text:
            continue

        # Linea separadora
        if set(text.strip()) == set('-') or set(text.strip()) == set('\u2014') or '---' in text:
            pdf.ln(2)
            pdf.set_draw_color(*GRIS_CLARO)
            y = pdf.get_y()
            pdf.line(margen_izq, y, page_w + margen_izq, y)
            pdf.ln(4)
            continue

        # Headings
        if 'Heading' in style_name or 'heading' in style_name:
            if '1' in style_name:
                pdf.set_font('Custom', 'B', 16)
                pdf.set_text_color(*AZUL_OSCURO)
                pdf.cell(page_w, 10, text, new_x='LMARGIN', new_y='NEXT')
                pdf.ln(4)
            elif '2' in style_name:
                pdf.set_font('Custom', 'B', 12)
                pdf.set_text_color(*AZUL_MEDIO)
                pdf.cell(page_w, 8, text, new_x='LMARGIN', new_y='NEXT')
                pdf.ln(3)
            else:
                pdf.set_font('Custom', 'B', 11)
                pdf.set_text_color(*AZUL_OSCURO)
                pdf.cell(page_w, 7, text, new_x='LMARGIN', new_y='NEXT')
                pdf.ln(2)
            continue

        # Items con bullet
        if text.startswith('- ') or text.startswith('* ') or text.startswith('\u2022 '):
            text_clean = text.lstrip('- *\u2022 ')
            text_clean = text_clean.strip()
            pdf.set_font('Custom', '', 10)
            pdf.set_text_color(*GRIS_TEXTO)
            x0 = pdf.get_x()
            pdf.cell(5, 5, '-')
            pdf.multi_cell(page_w - 5, 5, text_clean, new_x='LMARGIN', new_y='NEXT')
            pdf.ln(1)
            continue

        # Texto normal
        pdf.set_font('Custom', '', 10)
        pdf.set_text_color(*GRIS_TEXTO)
        pdf.multi_cell(page_w, 5, text)
        pdf.ln(2)

    # ── Tablas ──
    for table in doc.tables:
        pdf.ln(2)
        rows_data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)

        if not rows_data:
            continue

        n_cols = len(rows_data[0])
        col_w = page_w / n_cols

        # Header
        pdf.set_font('Custom', 'B', 7.5)
        pdf.set_fill_color(*AZUL_OSCURO)
        pdf.set_text_color(255, 255, 255)
        pdf.set_draw_color(*AZUL_OSCURO)
        for cell_text in rows_data[0]:
            pdf.cell(col_w, 6, cell_text, border=1, fill=True, align='C')
        pdf.ln()

        # Filas
        pdf.set_font('Custom', '', 7.5)
        pdf.set_text_color(*GRIS_TEXTO)
        pdf.set_draw_color(0xCC, 0xCC, 0xCC)

        for row_idx, row in enumerate(rows_data[1:], 1):
            if row_idx % 2 == 0:
                pdf.set_fill_color(245, 245, 245)
                fill = True
            else:
                fill = False

            max_lines = 1
            for cell_text in row:
                lines = pdf.multi_cell(col_w, 5, cell_text, split_only=True)
                max_lines = max(max_lines, len(lines))

            row_h = max(5.5, max_lines * 5)

            for cell_text in row:
                pdf.cell(col_w, row_h, cell_text, border=1, fill=fill, align='C')
            pdf.ln()

        pdf.ln(5)

    # ── Pie final ──
    pdf.ln(10)
    pdf.set_draw_color(*GRIS_CLARO)
    y = pdf.get_y()
    pdf.line(margen_izq, y, page_w + margen_izq, y)
    pdf.ln(5)
    pdf.set_font('Custom', 'I', 8)
    pdf.set_text_color(0x99, 0x99, 0x99)
    pdf.cell(page_w, 5, 'Informe generado automaticamente por Revenue Management Engine', align='C', new_x='LMARGIN', new_y='NEXT')
    pdf.cell(page_w, 5, 'Datos sinteticos 2025 - seed=42', align='C', new_x='LMARGIN', new_y='NEXT')

    pdf.output(str(pdf_path))
    return pdf_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Uso: python convertir_docx_a_pdf.py <input.docx> [output.pdf]")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(input_path):
        print(f"ERROR: No se encuentra el archivo: {input_path}")
        sys.exit(1)

    try:
        result = docx_a_pdf(input_path, output_path)
        print(f"OK PDF generado: {result}")
    except Exception as e:
        print(f"ERROR durante la conversion: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
